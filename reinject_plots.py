"""
Re-inject 7 result plots into ProjectNeal_IEEE_Final_Print.docx.
Walks all body elements including table cells to find anchors.
"""
import sys, io, copy, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = r'D:\Thrust Vector Control 2026\ProjectNeal_IEEE_Final_Print.docx'
DEST = r'D:\Thrust Vector Control 2026\ProjectNeal_IEEE_Final_Print.docx'
PLOTS = r'D:\Thrust Vector Control 2026\gen_plots'

doc = Document(SRC)

# ── collect ALL paragraphs in document order (body + tables) ──────────────────
from docx.oxml.ns import nsmap
from lxml import etree

def iter_block_items(doc_obj):
    """Yield paragraphs and tables in document order."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    body = doc_obj.element.body
    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            yield Paragraph(child, doc_obj)
        elif tag == 'tbl':
            yield Table(child, doc_obj)

def all_paragraphs(doc_obj):
    """Return flat list of all Paragraph objects in document order."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    result = []
    for item in iter_block_items(doc_obj):
        if isinstance(item, Paragraph):
            result.append(item)
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    result.extend(cell.paragraphs)
    return result

def has_image(para):
    xml = para._element.xml
    return 'blipFill' in xml or 'graphicData' in xml

def replace_image_in_para(para, img_path, width_in):
    p_elem = para._element
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'pPr':
            p_elem.remove(child)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def insert_after(anchor_para, img_path, width_in):
    """Insert image paragraph immediately after anchor_para."""
    from docx import Document as _D
    tmp = _D(); tmp_p = tmp.add_paragraph()
    run = tmp_p.add_run(); run.add_picture(img_path, width=Inches(width_in))
    run_elem = copy.deepcopy(tmp_p.runs[0]._element)

    new_p = OxmlElement('w:p')
    pPr   = OxmlElement('w:pPr')
    jc    = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc); new_p.append(pPr); new_p.append(run_elem)
    anchor_para._element.addnext(new_p)

# ── Build flat paragraph list ──────────────────────────────────────────────────
all_paras = all_paragraphs(doc)
print(f'Total paragraphs (incl. tables): {len(all_paras)}')

# Debug: print all para texts with index
for i, p in enumerate(all_paras):
    if p.text.strip():
        print(f'  [{i:4d}] {p.text.strip()[:100]}')

# ── Injections config ─────────────────────────────────────────────────────────
INJECTIONS = [
    ('State weighting matrix Q = diag',        'fig_all_controllers_overlay.png', 6.2),
    ('Fig. 14. Orientation output',             'fig_controller_comparison.png',   6.2),
    ('z3-dot = -beta3',                         'fig_adrc_eso.png',                6.0),
    ('Active Disturbance Rejection Control',    'fig_mrac_adaptation.png',         6.0),
    ('Digital Twin Controller Integration',     'fig_6dof_results.png',            6.2),
    ('Results and Discussion',                  'fig_monte_carlo.png',             6.2),
    ('LQI achieves the lowest mean pitch',      'fig_benchmark_bar.png',           6.2),
]

print('\n=== Injection pass ===')
for anchor_sub, plot_file, width in INJECTIONS:
    img_path = os.path.join(PLOTS, plot_file)
    if not os.path.exists(img_path):
        print(f'  [SKIP] file not found: {plot_file}')
        continue

    # Find anchor
    anchor_idx = None
    for i, p in enumerate(all_paras):
        if anchor_sub.lower() in p.text.lower():
            anchor_idx = i
            break

    if anchor_idx is None:
        print(f'  [MISS] anchor: "{anchor_sub[:55]}"')
        continue

    # Find next image para within 10 slots
    img_para = None
    for j in range(anchor_idx+1, min(anchor_idx+10, len(all_paras))):
        if has_image(all_paras[j]):
            img_para = all_paras[j]
            break

    if img_para:
        replace_image_in_para(img_para, img_path, width)
        print(f'  [REPL] {plot_file}  (after para {anchor_idx}: "{all_paras[anchor_idx].text[:50]}")')
    else:
        insert_after(all_paras[anchor_idx], img_path, width)
        print(f'  [INS]  {plot_file}  (after para {anchor_idx}: "{all_paras[anchor_idx].text[:50]}")')

doc.save(DEST)
print(f'\nSaved: {DEST}')
