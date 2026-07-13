"""
Surgical replacement: find FIRST image para after each anchor section,
replace only those two. All other images stay untouched.
Output: ProjectNeal_IEEE_Final_Print.docx
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SRC  = r'D:\Thrust Vector Control 2026\ProjectNeal_IEEE_WithPlots.docx'
DEST = r'D:\Thrust Vector Control 2026\ProjectNeal_IEEE_Final_Print.docx'
PLOTS = r'D:\Thrust Vector Control 2026\gen_plots'

doc = Document(SRC)
paras = doc.paragraphs

def has_image(para):
    xml = para._element.xml
    return 'blipFill' in xml or 'graphicData' in xml

def replace_image_para(para, img_path, width_in=6.5):
    """Wipe paragraph content and insert a centred image."""
    p_elem = para._element
    # remove all children except paragraph properties
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'pPr':
            p_elem.remove(child)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Diagnostic pass: print all paragraphs with images ─────────────────────────
print('=== Image paragraphs in document ===')
for i, p in enumerate(paras):
    if has_image(p):
        # also show the 3 paragraphs before for context
        ctx = ' | '.join(paras[max(0,i-3+j)].text.strip()[:60] for j in range(3))
        print(f'  [{i:3d}] context: {ctx}')

print()

# ── Find anchor paragraphs ─────────────────────────────────────────────────────
# Workflow: first image AFTER "IX" or "SYSTEM FLOW" section heading
# Simulink: first image AFTER "TABLE II" or "6DOF Model Parameters"

workflow_anchor_idx  = None
simulink_anchor_idx  = None

for i, p in enumerate(paras):
    t = p.text.strip().lower()
    if workflow_anchor_idx is None and ('system flow' in t or 'ix.' in t):
        workflow_anchor_idx = i
        print(f'Workflow anchor at [{i}]: "{p.text.strip()[:80]}"')
    if simulink_anchor_idx is None and ('table ii' in t or '6dof model param' in t):
        simulink_anchor_idx = i
        print(f'Simulink anchor at [{i}]: "{p.text.strip()[:80]}"')

def first_image_after(anchor_idx, paras, window=20):
    """Return index of first image paragraph after anchor."""
    if anchor_idx is None:
        return None
    for j in range(anchor_idx+1, min(anchor_idx+window, len(paras))):
        if has_image(paras[j]):
            return j
    return None

wf_img_idx  = first_image_after(workflow_anchor_idx, paras)
sim_img_idx = first_image_after(simulink_anchor_idx, paras)

print(f'\nWill replace:')
print(f'  Workflow  → para [{wf_img_idx}]')
print(f'  Simulink  → para [{sim_img_idx}]')

if wf_img_idx is not None:
    replace_image_para(paras[wf_img_idx],
                       os.path.join(PLOTS, 'fig_workflow_print.png'), 6.5)
    print(f'  [OK] Workflow image replaced.')

if sim_img_idx is not None:
    replace_image_para(paras[sim_img_idx],
                       os.path.join(PLOTS, 'fig_simulink_print.png'), 6.5)
    print(f'  [OK] Simulink image replaced.')

doc.save(DEST)
print(f'\nSaved: {DEST}')
